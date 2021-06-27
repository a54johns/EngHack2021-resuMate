import sys
import PyPDF2
import docx
import Levenshtein as lev
import string
from nltk import text
from nltk.corpus import wordnet
from PyQt5 import QtGui
from PyQt5.QtCore import QFileInfo
from PyQt5.QtGui import QStaticText, QIcon
from PyQt5.QtWidgets import QApplication, QMainWindow, QLabel, QPlainTextEdit, QTextEdit, QFileDialog, QPushButton, QComboBox, QGridLayout


class MainWindow(QMainWindow):
    globalFileName = None
    def __init__(self):
        #set up actual window
        super().__init__()
        self.setWindowTitle("ResuMate")
        self.setStyleSheet("QMainWindow {background: 'white';}")
        self.icon = QIcon('/Users/clairethompson/Desktop/EngHacks2021/Logo.png')
        self.setWindowIcon(self.icon)
        self.height = 800
        self.width = 800
        self.left = 10
        self.top = 10
        self.setGeometry(self.left, self.top,self.width, self.height)
        

        #set up widgets
        self.textbox = QTextEdit(self)
        self.textbox.move(350, 80)
        self.textbox.resize(400,220)
        self.textbox.verticalScrollBar().minimum()
        self.resultsbox = QTextEdit(self)
        self.button1 = QPushButton("Choose File", self)
        self.button1.clicked.connect(self.getOpenFileName)
        self.button2 = QPushButton("Load New Point", self)
        self.button2.clicked.connect(self.loadText)
        self.button3 = QPushButton("Generate Resume", self)
        self.button3.clicked.connect(self.backEnd)
        self.button3.setEnabled(False)
        self.button4 = QPushButton("Clear", self)
        self.button4.hide()
        self.button4.clicked.connect(self.clear)
        self.label1 = QLabel("Load Job Posting",self)
        self.label2 = QLabel("Add a new qualification point",self)
        self.label3 = QLabel("Loaded Job Posting:",self)
        self.label4 = QLabel(self)
        self.label5=QLabel(self)
        self.label6 = QLabel("Select prefered number of qualification points", self)
        self.comboBox = QComboBox(self)
        self.comboBox.addItems(['1','2','3','4','5'])
        self.comboBox.currentIndexChanged.connect(self.selectionChange)


        #arrange widgets
        self.button1.move(30,80)
        self.button2.move(349,320)
        self.button2.adjustSize() 
        self.button3.move(30, 325) 
        self.button3.adjustSize()
        self.label1.move(30,40)
        self.label1.adjustSize()
        self.label2.move(350, 40)
        self.label2.adjustSize()
        self.label3.move(30,140)
        self.label3.adjustSize()
        self.label6.move(30, 250)
        self.label6.adjustSize()
        self.comboBox.move(25, 270)

        #Style Widgets
        self.label1.setStyleSheet('color: #229396;')
        self.label2.setStyleSheet('color: #229396;')
        self.label3.setStyleSheet('color: #229396;')
        self.label4.setStyleSheet('color: #229396;')
        self.label5.setStyleSheet('color: #229396;')
        self.label6.setStyleSheet('color: #229396;')
        self.button1.setStyleSheet("background-color: #adb7b7; color: #229396;")
        self.button2.setStyleSheet("background-color: #adb7b7; color: #229396;")
        self.button3.setStyleSheet("background-color: #adb7b7; color: #229396;")

    #Gets job description pdf and displays file name once loaded
    def getOpenFileName(self):
        options = QFileDialog.Options()
        options |= QFileDialog.DontUseNativeDialog
        fileName, _ = QFileDialog.getOpenFileName(self,'Open file','PDF files (*.pdf)', options = options)
        if fileName:
            filename2 = QFileInfo(fileName).fileName()
            self.label4.setText(filename2) 
            self.label4.move(30,160)
            self.label4.setFixedWidth(300)
            self.globalFileName = fileName
        if self.globalFileName:
           self.button3.setEnabled(True)
            
    #Saves text from textbox to variable 
    def loadText(self):
        textValue = self.textbox.toPlainText()
        if textValue == "":
            self.label5.setText("No points entered")
        else: 
            self.label5.setText("Loaded Point: " + textValue)
            textValue
            document = docx.Document('/Users/clairethompson/Desktop/EngHacks2021/ResumePoints.docx')
            if textValue:
                document.add_paragraph(textValue)
        self.label5.move(355,355)
        self.label5.adjustSize()
        self.textbox.clear()   

    def selectionChange(self):
            self.num_items = self.comboBox.currentText()
            self.num_items = int(self.num_items)
            return  self.num_items
    
    def backEnd(self):
        pdfFileObj = open(self.globalFileName, 'rb')
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
        pageObj = pdfReader.getPage(0)
        pageObj1 = pdfReader.getPage(1)
        text = (pageObj.extractText()) + (pageObj1.extractText())
        post = text.split('Required Skills')
        Require = post[1].split('Application Information')
        Require = Require[0].split('Transportation and Housing')
        numberOfPoints = self.selectionChange()
        doc = docx.Document('/Users/clairethompson/Desktop/EngHacks2021/ResumePoints.docx')
        pa = doc.paragraphs
        point = []
        for i in range(len(pa)):
            if (i%2 ==0):
                 point.append(pa[i].text)

        keyPoints = self.textClean(point)
        Require = self.textClean(Require) 
        sim = []
        bestScore = 0
        newScore = 0
        sim = [0]*len(keyPoints)

        for i in range(len(keyPoints)):

            #split the resume point into the keywords
            sentence = keyPoints[i].split()

            for u in range(len(sentence)):
                word = sentence[u]
                bestWord = word
                bestScore = 0

                #for each word find common synonyms and compare to find similarities between resume points and posting
                for synonymSet in wordnet.synsets(word):
                    for lemma in synonymSet.lemmas():

                        #remake the sentence with the current synonyms 
                        tempSentence = " ".join(sentence)
                        newTempSentence = tempSentence.replace(word,lemma.name())

                        #compare the point against the job posting and develop a score
                        newScore = compareSentences(newTempSentence, Require[0].lower())

                        #Record the best score for the point for that synonym 
                        if newScore > bestScore:
                            bestScore = newScore
                            bestWord = lemma.name()

                #change the sentence to use the best synonym thus accumulating the overall best score for the point 
                sentence[u]=bestWord

                #resord the top score for that resume point 
            if bestScore > sim[i]:
                sim[i] = bestScore

        #sorting the sim scores 
        topPoints = sorted(range(len(sim)), key=lambda i: sim[i], reverse=True)[:numberOfPoints]

        #output the number of points (as requested by the user) and their unique ocrrelation score
        for k in range(len(topPoints)):
            #This is where you change the bottom text box
            self.resultsbox.insertPlainText(f'{sim[topPoints[k]]*100:2.2f} ' + point[topPoints[k]] + '\n')
            self.resultsbox.move(100,400)
            self.resultsbox.resize(550,300)
            self.resultsbox.setReadOnly(True)
            self.resultsbox.verticalScrollBar().minimum()
            self.button4.move(100, 720) 
            self.button4.show()

            #print(f'{sim[topPoints[k]]*100:2.2f}')
            #print(point[topPoints[k]])

        pdfFileObj.close()  

      #Function to clean code (remove punctuation, remove filler words, correct common pdf reader errors)
    def textClean (point):
        keyP = []

        #current list of common filler words used in job postings and resume points
        filler = ['you','this','good','to', 'and', 'a', 'for', 'with', 'by', 'in', 'the', 'as', 'such', 'of', 'experience', 'through', 'competencies', 'attained','completion', 'would','further', 'implemented', 
            'ensure', 'executing', 'at', 'previous', 'based','are', 'an', 'be', 'or', 'will', 'is', 'both', 'but', 'not', 'should', 'have', 'proficiency', 'related', 'taking', 'required', 'applicant', 'applicants',
            '4-month', '8-month', 'one', 'two', 'three', 'pursuing', 'ability', 'familiarity', 'knowledge', 'mandatory', 'some', 'understanding', 'quickly','come', 'considered', 'concepts', 'preferably', 'learn',
            'thorough', 'about', 'your', 'sense', 'placements', 'sept', 'date', '2021', 'discipline']

        #prepare punctuation library to required elements
        remov = string.punctuation
        remov = remov.replace("-", "")
        remov = remov.replace("/", "")

        for i in range(len(point)):

            #break the input into individual words
            words = point[i].split()
            for p in range(len(words)):

                #make all the words lowercase
                words[p] = words[p].lower()

                #remove all punctuation
                table = str.maketrans(dict.fromkeys(remov))
                words[p] = words[p].translate(table)

                #Remove characters requiring special care including pdf reader characters
                words[p]= words[p].split("/", 1)[0]
                words[p]= words[p].replace("Ł", " ")
                words[p]= words[p].replace("ł", " ")
                words[p]= words[p].replace("-€€", " ")
                words[p]= words[p].replace("€€", " ")
                words[p]= words[p].replace("  ", " ")

            #remove multi-word words when line breaks confuse the pdf reader
            for p in range(len(words)):
                if ' ' in words[p]:
                    firstHalf = (words[p].partition(' '))[0]
                    words[p] = firstHalf

            #sort words and remove duplicates as not to weight the scoring
            words.sort()
            finalList = list()

            for oneWord in words:
                if oneWord not in finalList:           
                    finalList.append(oneWord) 

            #Remove all known filler words
            for j in range(len(filler)):
                while (filler[j] in finalList):
                    finalList.remove(filler[j])

            #combine the remianing words
            text = " ".join(finalList)
            keyP.append(text)
            words.clear()

        #return the list of "cleaned" phrases
        return keyP

    #Function to perform comparison of resume points and selected job posting
    def compareSentences (app, point):

        #reset score
        score = 0

        #split sentences into array of words
        appSet = app.split(' ')
        pointSet = point.split(' ')

        #resets 
        bestConnectScore = 0
        connectScore = 0
        scoreArray = [0]*len(pointSet)
        refinedArray = []

        for i in range(len(pointSet)):
            bestConnectScore = 0
            connectScore = 0

            for j in range(len(appSet)):

                #apply Levenshtein theorm of comparison to determine the similarities in keywords
                connectScore = lev.ratio(pointSet[i], appSet[j])
                if connectScore > bestConnectScore:
                    bestConnectScore = connectScore

            #store the best score
            scoreArray[i] = bestConnectScore

        #apply a weighting function - only consider and record scores over 80% similar
        for i in range(len(scoreArray)):
            if scoreArray[i] > 0.5:
                refinedArray.append(scoreArray[i])

        #average the scores to determine an overall for the resume point 
        if len(refinedArray) != 0:
            score = sum(refinedArray) / len(refinedArray)
        else:
            score = 0

        return score

    def clear(self):
        self.resultsbox.clear()

if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())

